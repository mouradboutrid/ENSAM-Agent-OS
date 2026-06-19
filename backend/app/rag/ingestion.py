from __future__ import annotations
import csv
import io
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional

from langchain_text_splitters import RecursiveCharacterTextSplitter


@dataclass
class ParsedDocument:
    content: str
    metadata: dict = field(default_factory=dict)
    filename: str = ""
    file_type: str = ""
    page_count: int = 0


@dataclass
class DocumentChunk:
    text: str
    metadata: dict = field(default_factory=dict)
    chunk_index: int = 0
    token_estimate: int = 0


class FileParser:
    @staticmethod
    def parse(file_path: str, content_bytes: Optional[bytes] = None) -> ParsedDocument:
        path = Path(file_path)
        suffix = path.suffix.lower()
        parsers = {
            ".pdf": FileParser._parse_pdf,
            ".docx": FileParser._parse_docx,
            ".md": FileParser._parse_text,
            ".txt": FileParser._parse_text,
            ".csv": FileParser._parse_csv,
        }
        parser = parsers.get(suffix)
        if not parser:
            raise ValueError(f"Unsupported file type: {suffix}")
        doc = parser(file_path, content_bytes)
        doc.filename = path.name
        doc.file_type = suffix
        doc.metadata["source"] = path.name
        doc.metadata["file_type"] = suffix
        return doc

    @staticmethod
    def _parse_pdf(file_path: str, content_bytes: Optional[bytes] = None) -> ParsedDocument:
        from pypdf import PdfReader
        if content_bytes:
            reader = PdfReader(io.BytesIO(content_bytes))
        else:
            reader = PdfReader(file_path)
        pages = []
        for i, page in enumerate(reader.pages):
            text = page.extract_text() or ""
            pages.append(text)
        return ParsedDocument(
            content="\n\n".join(pages),
            page_count=len(pages),
            metadata={"total_pages": len(pages)},
        )

    @staticmethod
    def _parse_docx(file_path: str, content_bytes: Optional[bytes] = None) -> ParsedDocument:
        from docx import Document
        if content_bytes:
            doc = Document(io.BytesIO(content_bytes))
        else:
            doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        return ParsedDocument(content="\n\n".join(paragraphs), metadata={"paragraphs": len(paragraphs)})

    @staticmethod
    def _parse_text(file_path: str, content_bytes: Optional[bytes] = None) -> ParsedDocument:
        if content_bytes:
            content = content_bytes.decode("utf-8", errors="replace")
        else:
            content = Path(file_path).read_text(encoding="utf-8", errors="replace")
        return ParsedDocument(content=content)

    @staticmethod
    def _parse_csv(file_path: str, content_bytes: Optional[bytes] = None) -> ParsedDocument:
        if content_bytes:
            text = content_bytes.decode("utf-8", errors="replace")
        else:
            text = Path(file_path).read_text(encoding="utf-8", errors="replace")
        reader = csv.reader(io.StringIO(text))
        rows = list(reader)
        formatted = "\n".join([" | ".join(row) for row in rows])
        return ParsedDocument(content=formatted, metadata={"rows": len(rows)})


class TextChunker:
    def __init__(self, chunk_size: int = 512, chunk_overlap: int = 50):
        self._splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", ". ", " ", ""],
        )

    def chunk(self, document: ParsedDocument) -> list[DocumentChunk]:
        texts = self._splitter.split_text(document.content)
        chunks = []
        for i, text in enumerate(texts):
            chunk = DocumentChunk(
                text=text,
                metadata={**document.metadata, "chunk_index": i, "total_chunks": len(texts)},
                chunk_index=i,
                token_estimate=int(len(text.split()) * 1.3),
            )
            chunks.append(chunk)
        return chunks
